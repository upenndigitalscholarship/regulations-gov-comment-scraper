import os
import requests
import json
import time
from docx import Document
import urllib

header = "https://api.data.gov/regulations/v3/document.json?api_key=sILE0BdGLfTiMoa0JG7mssuvpPHC06bNK72u9As6&documentId="
apiKey = #add api key here"


documentIDs = [#add document ids here]

os.chdir(#add target directory here)

for document in documentIDs:
    thisDocument = header + document
    comment = json.loads(requests.get(thisDocument).text)
    submitterName = comment['submitterName']['value']
    organization = comment['organization']['value']
    fullComment = comment['comment']['value']
    attachments = comment['attachments']

    i = 1
    for a in attachments:
        sat = a['fileFormats'][0]
        strarray = str(sat).split('?', 2)
        attachurl = str(strarray[0]) + '?' + apiKey + '&' + strarray[1]
        attachf = document + "_(" + str(i) + ")_" + "Attachment.docx"
        urllib.request.urlretrieve(attachurl, attachf)
        i = i + 1

    commentDoc = Document();

    commentDoc.add_paragraph(submitterName)
    commentDoc.add_paragraph(organization)
    commentDoc.add_paragraph(fullComment)

    commentDoc.save(document + '.docx')
    print("reached document " + document)

    time.sleep(3)
